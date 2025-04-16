import time
from datetime import datetime
from google.oauth2 import service_account
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from googleapiclient.http import MediaFileUpload
from docx import Document
import os
import io
from googleapiclient.http import MediaIoBaseDownload

SPREADSHEET_ID = "1MhK40R1-GB4h8oTgq3624x7sLFCa7AD6TKLHlnWJDlU"

SHEET_SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive"]
CURRENCY_TEXT_SYMBOLS = {
    "usd": "$",
    "eur": "€",
    "gbp": "£",
    "aud": "A$",
    "zar": "R"
}

REGIONS_DATA = {
    "South Asia": {"name": "South Asia", "countries": "India"},
    "Europe* (Ukraine)": {"name": "Europe*", "countries": "Ukraine"},
    "Central Asia": {"name": "Central Asia", "countries": "Turkey, Kazakhstan, Kyrgyzstan"},
    "Latam": {"name": "Latam", "countries": "Mexico, Argentina, Colombia, Brazil"},
    "Europe": {"name": "Europe", "countries": "Poland, Portugal, Spain, Czech Republic, Latvia, Greece, Georgia, Slovak Republic, Moldova, Lithuania, etc."}
}

TEMPLATES = {
    "1_reg": "1FyGrmUQMr3XWXNcHoeSjFHexjfeTd_ni",
    "2_reg": "1FBSMDWt6NC3qo4EDks9ghgRMyQRuxtqO",
    "3_reg": "1qeJDe7wx3aem6az9qlwFipbPL7SrRdCQ",
    "4_reg": "1BVR6oMyw92ytJCPNCMIJ2b0C3w-euKIO",
    "5_reg": "1uBNzVbwrqyRgPZh875S0tE75D17TTx0K"
}

class RateCard:
    def __init__(self, currency):
        self.currency = currency
        self.rates = []
        self.regions = []
        self.drive_service = None
        self.rates_by_title_region = {}
        self.file_name = ""
        self.shared_drive_id = "1-VU-OtThHDrgyMCQlhYiLI-CqA2gGREb"
        self.rate_card_drive_id = None

    def init_drive(self):
        creds = Credentials.from_authorized_user_file("tokens/drive.json", DRIVE_SCOPES)
        creds.refresh(Request())
        self.drive_service = build("drive", "v3", credentials=creds)

    def download_file(self, file_id, file_name):
        request = self.drive_service.files().get_media(fileId=file_id)
        file = io.BytesIO()
        downloader = MediaIoBaseDownload(file, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            print(f"Template Download {int(status.progress() * 100)}% complete.")
        with open(file_name, 'wb') as f:
            f.write(file.getvalue())
        print(f"Template '{file_name}' downloaded successfully.")

    def collect_inputs(self):
        creds = service_account.Credentials.from_service_account_file(
            'tokens/sheet.json', scopes=SHEET_SCOPES
        )
        service = build('sheets', 'v4', credentials=creds)
        sheet = service.spreadsheets()
        # self.currency = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range='Original External Rates (currency)!O2').execute()['values'][0][0]
        if self.currency == "usd":
            rates_rage = 'Original External Rates in USD !A:H'
            # self.regions = ['Europe* (Ukraine)', 'Latam', 'Europe', 'Central Asia', 'South Asia']
            self.regions = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range='Original External Rates in USD !J5').execute()[
                'values'][0][0].split(", ")
        else:
            rates_rage = 'Original External Rates (currency)!A:H'
            self.regions = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range='Original External Rates (currency)!J5').execute()[
                'values'][0][0].split(", ")
        print(self.regions)
        rates = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=rates_rage).execute()['values']
        for rate in rates:
            if not rate:
                continue
            if rate[0]:
                self.rates.append(rate)

    def prepare_rates(self):
        header_row = self.rates[0]
        counter = 0
        regions = {
            header_row[3]: 3,
            header_row[4]: 4,
            header_row[5]: 5,
            header_row[6]: 6,
            header_row[7]: 7

        }
        rates_by_title = {}

        for rate_row in self.rates:
            counter += 1
            if counter == 1:
                continue
            title = rate_row[1].replace('Engeneer', 'Engineer').strip()
            for region, rate_index in regions.items():
                min_rate = rate_row[rate_index].replace("$", "")
                max_rate = rate_row[rate_index].replace("$", "")
                if min_rate == max_rate:
                    if not min_rate:
                        rate_str = ""
                    elif self.currency in ['eur', "gbp", "usd"]:
                        rate_str = f"{CURRENCY_TEXT_SYMBOLS[self.currency]}{min_rate}"
                    else:
                        rate_str = f"{min_rate} {self.currency.upper()}"
                else:
                    rate_str = f"{CURRENCY_TEXT_SYMBOLS[self.currency]}{min_rate}-{CURRENCY_TEXT_SYMBOLS[self.currency]}{max_rate}/h"
                rates_by_title[f"{title}_{region}"] = rate_str
        self.rates_by_title_region = rates_by_title

    def replace_rates_in_doc(self):
        print(f"Generating CV for {len(self.regions)} regions")
        template_id = TEMPLATES[f'{len(self.regions)}_reg']
        self.download_file(template_id, 'template/template.docx')
        doc = Document(f'template/template.docx')
        all_regions = list(REGIONS_DATA.keys())
        regions_to_set = []
        countries_to_set = []
        for region in all_regions:
            if region in self.regions:
                regions_to_set.append(region)
        for region in regions_to_set:
            countries_to_set.append(REGIONS_DATA[region]['countries'])

        for table in doc.tables:
            regions_row = table.rows[1]
            countries_row = table.rows[2]
            for cell in regions_row.cells:
                if 'region_' in cell.text:
                    region_index = int(cell.text.replace("region_", "")) - 1
                    region_name = REGIONS_DATA[regions_to_set[region_index]]['name']
                    cell.paragraphs[0].runs[0].text = region_name
            for cell in countries_row.cells:
                if 'region_' in cell.text:
                    region_index = int(cell.text.replace("region_", "").replace("_countries", "")) - 1
                    country_names = countries_to_set[region_index]
                    cell.paragraphs[0].runs[0].text = country_names

        # fill rates
        for table in doc.tables:
            row_counter = 0
            rate_headers = []
            for row in table.rows:
                row_counter += 1
                if row_counter == 2:
                    rate_headers = []
                    start_index = 1
                    for region in regions_to_set:
                        start_index += 2
                        rate_headers.append(row.cells[start_index])
                if row_counter < 5:
                    continue
                title_cell = row.cells[1]
                rates = []
                start_index = 1
                for region in regions_to_set:
                    start_index += 2
                    rates.append(row.cells[start_index])
                if not title_cell.text:
                    continue
                counter = 0
                for rate in rates:
                    rate_header = rate_headers[counter]
                    region = rate_header.text
                    if region == "Europe*":
                        region = "Europe* (Ukraine)"
                    rate_to_set = self.rates_by_title_region[f"{title_cell.text.strip()}_{region}"]
                    rate.paragraphs[0].runs[0].text = rate_to_set
                    counter += 1
        self.file_name = f"rate_card_{self.currency.lower()}_{datetime.today().strftime('%Y-%m-%d_%H:%M:%S')}.docx"
        doc.save(f"output/{self.file_name}")

    def upload_to_drive(self):
        file_metadata = {
            "name": self.file_name,
            "mimeType": "application/vnd.google-apps.document",
            "parents": [self.shared_drive_id]
        }
        media = MediaFileUpload(f"output/{self.file_name}")
        create_file = self.drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id",
            supportsAllDrives=True
        ).execute()
        self.rate_card_drive_id = create_file['id']

    def clear_cache(self):
        folder_paths = ['output', 'template']
        for folder_path in folder_paths:
            for file_name in os.listdir(folder_path):
                file_path = os.path.join(folder_path, file_name)
                if os.path.isfile(file_path):
                    os.remove(file_path)

    def generate_card(self):
        self.init_drive()
        self.collect_inputs()
        self.prepare_rates()
        self.replace_rates_in_doc()
        self.upload_to_drive()
        self.clear_cache()


def rate_card_generator(currency):
    rate_card_handler = RateCard(currency)
    try:
        rate_card_handler.generate_card()
    except Exception as e:
        print(e)
        pass
    return rate_card_handler.rate_card_drive_id


